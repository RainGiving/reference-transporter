from __future__ import annotations

import json
import re
import time
import uuid
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile

import requests
from docx import Document
from lxml import etree

from .metadata_resolver import MetadataResolution, MetadataResolver, extract_identifiers
from .utils import clean_whitespace, normalize_doi, random_id
from .word import WordZoteroInjector

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
CP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
VT_NS = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"

CITATION_TOKEN_RE = re.compile(r"\[(\d+(?:\s*[-–,，]\s*\d+)*)\]")
REFERENCE_RE = re.compile(r"^\[(\d+)\]\s*(.+?)\s*$")
TYPE_RE = re.compile(r"^(?P<authors>.+?)\.\s+(?P<title>.+?)\[(?P<rtype>[A-Z/]+)\](?P<rest>.*)$")
INITIAL_TOKEN_RE = re.compile(r"^[A-Z](?:-[A-Z])?$")
TITLE_NORMALIZER_RE = re.compile(r"[^a-z0-9]+")


@dataclass(slots=True)
class ParsedReference:
    number: int
    raw: str
    authors_raw: str
    title: str
    ref_type_code: str
    item_type: str
    creators: list[dict]
    fields: dict
    confidence: str = "parsed"
    item_key: str | None = None
    resolution_source: str | None = None
    resolution_score: float | None = None
    used_fallback: bool = False
    failure_reason: str | None = None
    created_new_item: bool = False


def normalize_title(value: str) -> str:
    return TITLE_NORMALIZER_RE.sub("", value.lower())


def normalize_heading(value: str) -> str:
    return value.replace("\u200b", "").replace("\xa0", " ").strip()


def find_reference_heading_index(doc: Document) -> int:
    for index, para in enumerate(doc.paragraphs):
        if normalize_heading(para.text) == "参考文献":
            return index
    raise ValueError("Could not find '参考文献' heading in document")


def expand_citation_numbers(token: str) -> list[int]:
    token = token.strip()[1:-1]
    values: list[int] = []
    for part in re.split(r"[，,]", token):
        part = part.strip()
        if not part:
            continue
        if "-" in part or "–" in part:
            left, right = re.split(r"[-–]", part, maxsplit=1)
            start = int(left.strip())
            end = int(right.strip())
            step = 1 if end >= start else -1
            values.extend(list(range(start, end + step, step)))
        else:
            values.append(int(part))
    return values


def extract_reference_paragraphs(docx_path: str | Path) -> list[str]:
    doc = Document(str(docx_path))
    start = find_reference_heading_index(doc) + 1
    refs: list[str] = []
    for para in doc.paragraphs[start:]:
        text = normalize_heading(para.text)
        if not text:
            continue
        if REFERENCE_RE.match(text):
            refs.append(text)
    return refs


def load_references_from_text(path: str | Path) -> list[str]:
    content = Path(path).read_text(encoding="utf-8")
    raw_blocks = [clean_whitespace(block.replace("\n", " ")) for block in re.split(r"\n\s*\n", content) if clean_whitespace(block.replace("\n", " "))]
    if not raw_blocks:
        return []

    if len(raw_blocks) == 1:
        lines = [normalize_heading(line) for line in content.splitlines() if normalize_heading(line)]
        if all(REFERENCE_RE.match(line) for line in lines):
            return lines
        if all(not REFERENCE_RE.match(line) for line in lines):
            return [f"[{idx}] {line}" for idx, line in enumerate(lines, start=1)]

    normalized = []
    for idx, block in enumerate(raw_blocks, start=1):
        normalized.append(block if REFERENCE_RE.match(block) else f"[{idx}] {block}")
    return normalized


def _extract_docx_zotero_document_data(files: dict[str, bytes]) -> dict | None:
    custom_xml = files.get("docProps/custom.xml")
    if not custom_xml:
        return None
    try:
        root = etree.fromstring(custom_xml)
    except Exception:
        return None

    chunks: list[tuple[int, str]] = []
    for prop in root.findall(f"{{{CP_NS}}}property"):
        name = prop.get("name", "")
        if not name.startswith("ZOTERO_PREF_"):
            continue
        try:
            order = int(name.split("_")[-1])
        except ValueError:
            order = 9999
        text = "".join(prop.xpath(".//vt:lpwstr/text()", namespaces={"vt": VT_NS}))
        chunks.append((order, text))
    if not chunks:
        return None
    payload = "".join(text for _, text in sorted(chunks))
    payload = payload.strip()
    if not payload:
        return None

    if payload.startswith("{"):
        try:
            return json.loads(payload)
        except Exception:
            return None
    if payload.startswith("<"):
        try:
            root = etree.fromstring(payload.encode("utf-8"))
            style = root.find("style")
            session = root.find("session")
            prefs = root.find("prefs")
            result = {
                "style": {
                    "styleID": style.get("id") if style is not None else "",
                    "locale": style.get("locale") if style is not None else "",
                },
                "sessionID": session.get("id") if session is not None else "",
                "prefs": {},
            }
            if prefs is not None:
                for pref in prefs.findall("pref"):
                    result["prefs"][pref.get("name")] = pref.get("value")
            return result
        except Exception:
            return None
    return None


def _find_zotero_prefs_js() -> Path | None:
    base = Path.home() / "AppData/Roaming/Zotero/Zotero/Profiles"
    if not base.exists():
        return None
    prefs_files = [profile / "prefs.js" for profile in base.iterdir() if (profile / "prefs.js").exists()]
    if not prefs_files:
        return None
    return max(prefs_files, key=lambda p: p.stat().st_mtime)


def _extract_current_zotero_style_preferences() -> tuple[str | None, str | None]:
    prefs_path = _find_zotero_prefs_js()
    if not prefs_path:
        return None, None
    text = prefs_path.read_text(encoding="utf-8", errors="ignore")
    style_match = re.search(r'user_pref\("extensions\.zotero\.export\.lastStyle",\s*"([^"]+)"\);', text)
    locale_match = re.search(r'user_pref\("extensions\.zotero\.export\.lastLocale",\s*"([^"]+)"\);', text)
    style_id = style_match.group(1) if style_match else None
    locale = locale_match.group(1) if locale_match else None
    return style_id, locale


def resolve_style_preferences(input_docx: str | Path, style_id: str | None = None, locale: str | None = None) -> tuple[str, str]:
    inherited_style = None
    inherited_locale = None
    with ZipFile(input_docx, "r") as source:
        files = {name: source.read(name) for name in source.namelist() if name in {"docProps/custom.xml"}}
    doc_data = _extract_docx_zotero_document_data(files)
    if doc_data:
        inherited_style = clean_whitespace((doc_data.get("style") or {}).get("styleID", ""))
        inherited_locale = clean_whitespace((doc_data.get("style") or {}).get("locale", ""))

    zotero_style, zotero_locale = _extract_current_zotero_style_preferences()
    resolved_style = style_id or inherited_style or zotero_style or "http://www.zotero.org/styles/apa"
    resolved_locale = locale or inherited_locale or zotero_locale or "en-US"
    return resolved_style, resolved_locale


def parse_reference(raw_reference: str) -> ParsedReference:
    outer = REFERENCE_RE.match(normalize_heading(raw_reference))
    if not outer:
        raise ValueError(f"Invalid reference paragraph: {raw_reference}")
    number = int(outer.group(1))
    body = outer.group(2).strip()
    body = body.replace("\xa0", " ")
    match = TYPE_RE.match(body)
    if not match:
        raise ValueError(f"Unsupported reference format: {raw_reference}")

    authors_raw = clean_whitespace(match.group("authors"))
    title = clean_whitespace(match.group("title"))
    ref_type_code = match.group("rtype")
    rest = clean_whitespace(match.group("rest").lstrip(". "))
    creators = parse_creators(authors_raw)

    if ref_type_code == "J":
        item_type = "journalArticle"
        fields = parse_journal_fields(rest)
    elif ref_type_code == "C":
        item_type = "conferencePaper"
        fields = parse_conference_fields(rest)
    elif ref_type_code == "EB/OL":
        item_type, fields = parse_online_fields(title, rest)
    else:
        raise ValueError(f"Unsupported reference type: {ref_type_code}")

    fields["title"] = title
    return ParsedReference(
        number=number,
        raw=raw_reference,
        authors_raw=authors_raw,
        title=title,
        ref_type_code=ref_type_code,
        item_type=item_type,
        creators=creators,
        fields=fields,
    )


def parse_creators(authors_raw: str) -> list[dict]:
    parts = [clean_whitespace(part) for part in authors_raw.split(",")]
    creators: list[dict] = []
    for part in parts:
        normalized = part.lower().replace(".", "").strip()
        if not part or normalized == "et al":
            continue
        tokens = part.split()
        initials: list[str] = []
        while tokens and INITIAL_TOKEN_RE.match(tokens[-1]):
            initials.insert(0, tokens.pop())
        if initials and tokens:
            creators.append(
                {
                    "creatorType": "author",
                    "firstName": " ".join(initials),
                    "lastName": " ".join(tokens),
                }
            )
        else:
            creators.append({"creatorType": "author", "name": part})
    return creators


def parse_journal_fields(rest: str) -> dict:
    match = re.match(r"^(?P<journal>.+?),\s*(?P<year>\d{4}),\s*(?P<after>.+?)\.?$", rest)
    if not match:
        raise ValueError(f"Could not parse journal reference tail: {rest}")
    journal = clean_whitespace(match.group("journal"))
    year = match.group("year")
    after = clean_whitespace(match.group("after"))
    volume = ""
    issue = ""
    pages = ""
    if ":" in after:
        vol_issue, pages = [clean_whitespace(x) for x in after.split(":", 1)]
    else:
        vol_issue = after
    issue_match = re.match(r"^(?P<vol>[^()]+?)(?:\((?P<issue>[^)]+)\))?$", vol_issue)
    if issue_match:
        volume = clean_whitespace(issue_match.group("vol"))
        issue = clean_whitespace(issue_match.group("issue"))
    else:
        volume = vol_issue
    return {
        "publicationTitle": journal,
        "date": year,
        "volume": volume,
        "issue": issue,
        "pages": pages.rstrip("."),
    }


def parse_conference_fields(rest: str) -> dict:
    tail = rest
    if tail.startswith("//"):
        tail = tail[2:].strip()
    if ". " in tail:
        book_title, meta = tail.split(". ", 1)
    else:
        book_title, meta = tail, ""
    book_title = book_title.rstrip(".")
    year_match = re.search(r",\s*(\d{4})(?:,\s*(.+))?$", meta)
    place = ""
    year = ""
    vol_issue_pages = ""
    if year_match:
        year = year_match.group(1)
        place = clean_whitespace(meta[: year_match.start()])
        vol_issue_pages = clean_whitespace(year_match.group(2))
    volume = ""
    pages = ""
    if vol_issue_pages:
        if ":" in vol_issue_pages:
            volume, pages = [clean_whitespace(x) for x in vol_issue_pages.split(":", 1)]
        else:
            volume = vol_issue_pages
    return {
        "proceedingsTitle": book_title,
        "conferenceName": book_title,
        "place": place.rstrip(","),
        "date": year,
        "volume": volume.rstrip("."),
        "pages": pages.rstrip("."),
    }


def parse_online_fields(title: str, rest: str) -> tuple[str, dict]:
    years = re.findall(r"(?:19|20)\d{2}", rest)
    year = years[-1] if years else ""
    url_match = re.search(r"https?://\S+", rest)
    url = url_match.group(0).rstrip(".,") if url_match else ""
    if "arxiv:" in rest.lower():
        archive_id = rest.split(":", 1)[1].split(",")[0].strip()
        return (
            "preprint",
            {
                "date": year,
                "url": f"https://arxiv.org/abs/{archive_id}",
                "archiveID": archive_id,
                "repository": "arXiv",
            },
        )
    website_title = ""
    if url:
        website_title = re.sub(r"^www\.", "", requests.utils.urlparse(url).netloc)
    return (
        "webpage",
        {
            "date": year,
            "url": url,
            "websiteTitle": website_title,
        },
    )


def maybe_enrich_with_crossref(parsed: ParsedReference, client: requests.Session) -> None:
    if parsed.item_type != "journalArticle":
        return
    query = f"{parsed.title}. {parsed.fields.get('publicationTitle', '')}. {parsed.fields.get('date', '')}."
    try:
        response = client.get(
            "https://api.crossref.org/works",
            params={"query.bibliographic": query, "rows": 5, "mailto": "yuqing@example.com"},
            timeout=20,
        )
        response.raise_for_status()
    except Exception:
        return

    best = None
    best_score = 0.0
    expected_title = normalize_title(parsed.title)
    expected_year = parsed.fields.get("date", "")
    expected_author = ""
    if parsed.creators:
        first = parsed.creators[0]
        expected_author = (first.get("lastName") or first.get("name") or "").lower()

    for item in response.json().get("message", {}).get("items", []):
        title = clean_whitespace((item.get("title") or [""])[0])
        title_score = SequenceMatcher(None, expected_title, normalize_title(title)).ratio()
        year = ""
        for key in ("published-print", "published-online", "issued", "created"):
            part = item.get(key, {}).get("date-parts")
            if part and part[0]:
                year = str(part[0][0])
                break
        year_score = 0.2 if year == expected_year else 0.0
        authors = item.get("author", [])
        author_score = 0.0
        if authors:
            family = authors[0].get("family", "").lower()
            if expected_author and family == expected_author:
                author_score = 0.1
        score = title_score + year_score + author_score
        if score > best_score:
            best = item
            best_score = score

    if best and best_score >= 1.0:
        doi = normalize_doi(best.get("DOI"))
        if doi:
            parsed.fields["DOI"] = doi
        url = best.get("URL")
        if url:
            parsed.fields["url"] = url
        parsed.confidence = "crossref"


def item_completeness_score(item_data: dict) -> float:
    score = 0.0
    for field in [
        "DOI",
        "url",
        "pages",
        "volume",
        "issue",
        "publicationTitle",
        "proceedingsTitle",
        "conferenceName",
        "place",
        "date",
        "archiveID",
        "repository",
        "journalAbbreviation",
        "ISBN",
    ]:
        if clean_whitespace(str(item_data.get(field, ""))):
            score += 1.0
    creators = item_data.get("creators", [])
    score += min(len(creators), 8) * 0.35
    return score


def build_item_payload(parsed: ParsedReference, resolved: MetadataResolution | None) -> dict:
    item = to_connector_item(parsed)
    if resolved:
        item["itemType"] = resolved.item.get("itemType", item.get("itemType", parsed.item_type))
        prefer_longer_text_fields = {"publicationTitle", "proceedingsTitle", "conferenceName", "place"}
        for key, value in resolved.item.items():
            if key == "creators":
                if value:
                    item[key] = value
                continue
            if value not in (None, "", []):
                if key in prefer_longer_text_fields:
                    current = clean_whitespace(str(item.get(key, "")))
                    candidate = clean_whitespace(str(value))
                    item[key] = candidate if len(candidate) > len(current) else current
                else:
                    item[key] = value

    item.setdefault("title", parsed.title)
    item.setdefault("creators", parsed.creators)
    item.setdefault("itemType", parsed.item_type)

    marker = str(uuid.uuid4())
    source = resolved.source if resolved else "fallback-parsed"
    ref_tag = f"zwb-ref-{parsed.number}"
    source_tag = f"zwb-source-{source}"
    existing_tags = [tag for tag in item.get("tags", []) if isinstance(tag, dict) and tag.get("tag")]
    item["tags"] = existing_tags + [{"tag": "zwb-import"}, {"tag": ref_tag}, {"tag": source_tag}]

    extra = clean_whitespace(item.get("extra", ""))
    extra_lines = [line for line in extra.splitlines() if line.strip()]
    extra_lines.extend(
        [
            f"ZWB Import ID: {marker}",
            f"ZWB Ref Number: {parsed.number}",
            f"ZWB Source: {source}",
        ]
    )
    item["extra"] = "\n".join(extra_lines)
    item["_zwb_import_id"] = marker
    return item


def to_connector_item(parsed: ParsedReference) -> dict:
    item = {
        "itemType": parsed.item_type,
        "title": parsed.fields["title"],
        "creators": parsed.creators,
    }
    for key, value in parsed.fields.items():
        if key == "title":
            continue
        if value:
            item[key] = value
    return item


class ZoteroConnectorClient:
    def __init__(self, base_url: str = "http://127.0.0.1:23119") -> None:
        self.base_url = base_url.rstrip("/")
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": "zotero-word-bridge/0.1"})

    def get_selected_collection(self) -> dict:
        response = self.session.post(f"{self.base_url}/connector/getSelectedCollection", json={}, timeout=20)
        response.raise_for_status()
        return response.json()

    def save_item(self, item: dict, target: str | None = None) -> str:
        session_id = str(uuid.uuid4())
        payload_item = {k: v for k, v in item.items() if not k.startswith("_zwb_")}
        payload = {"sessionID": session_id, "items": [payload_item]}
        response = self.session.post(f"{self.base_url}/connector/saveItems", json=payload, timeout=60)
        response.raise_for_status()
        if target:
            update = self.session.post(
                f"{self.base_url}/connector/updateSession",
                json={"sessionID": session_id, "target": target},
                timeout=60,
            )
            update.raise_for_status()
        return session_id


class ZoteroLocalCollection:
    def __init__(self, api_base_url: str = "http://127.0.0.1:23119/api") -> None:
        self.api_base_url = api_base_url.rstrip("/")
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": "zotero-word-bridge/0.1"})

    def list_collections(self) -> list[dict]:
        response = self.session.get(f"{self.api_base_url}/users/0/collections", params={"format": "json"}, timeout=30)
        response.raise_for_status()
        return response.json()

    def resolve_collection_key(self, name: str) -> str:
        matches = [entry["key"] for entry in self.list_collections() if entry["data"]["name"] == name]
        if not matches:
            raise ValueError(f"Collection not found: {name}")
        if len(matches) > 1:
            raise ValueError(f"Collection name is ambiguous: {name}")
        return matches[0]

    def list_collection_items(self, collection_key: str) -> list[dict]:
        response = self.session.get(
            f"{self.api_base_url}/users/0/collections/{collection_key}/items",
            params={"format": "json"},
            timeout=30,
        )
        response.raise_for_status()
        return response.json()

    def fetch_csljson(self, item_key: str) -> dict:
        response = self.session.get(
            f"{self.api_base_url}/users/0/items/{item_key}",
            params={"format": "csljson"},
            timeout=30,
        )
        response.raise_for_status()
        data = response.json()
        return data[0] if isinstance(data, list) else data

    @staticmethod
    def build_item_uri(item_key: str) -> str:
        return f"http://zotero.org/users/0/items/{item_key}"


def find_matching_item(parsed: ParsedReference, collection_items: Iterable[dict]) -> dict | None:
    target_title = normalize_title(parsed.title)
    target_year = parsed.fields.get("date", "")
    target_doi = normalize_doi(parsed.fields.get("DOI"))
    for item in collection_items:
        data = item["data"]
        if target_doi and normalize_doi(data.get("DOI")) == target_doi:
            return item
    for item in collection_items:
        data = item["data"]
        if normalize_title(data.get("title", "")) == target_title:
            existing_year = clean_whitespace(data.get("date", ""))
            if not target_year or not existing_year or target_year in existing_year:
                return item
    return None


def find_best_existing_item(candidate_item: dict, collection_items: Iterable[dict]) -> dict | None:
    doi = normalize_doi(candidate_item.get("DOI"))
    title = normalize_title(candidate_item.get("title", ""))
    exact_title_matches: list[dict] = []
    for item in collection_items:
        data = item["data"]
        if doi and normalize_doi(data.get("DOI")) == doi:
            return item
        if title and normalize_title(data.get("title", "")) == title:
            exact_title_matches.append(item)
    if not exact_title_matches:
        return None
    return max(exact_title_matches, key=lambda x: item_completeness_score(x["data"]))


def should_reuse_existing_item(existing_item: dict, candidate_item: dict, resolved: MetadataResolution | None) -> bool:
    if not resolved:
        return True
    existing_score = item_completeness_score(existing_item["data"])
    candidate_score = item_completeness_score(candidate_item)
    existing_doi = normalize_doi(existing_item["data"].get("DOI"))
    candidate_doi = normalize_doi(candidate_item.get("DOI"))
    if candidate_doi and existing_doi == candidate_doi:
        return True
    return existing_score >= candidate_score - 0.4


def import_references_to_collection(
    references: list[ParsedReference],
    collection_name: str = "master degree",
    connector_base_url: str = "http://127.0.0.1:23119",
    api_base_url: str = "http://127.0.0.1:23119/api",
) -> list[ParsedReference]:
    connector = ZoteroConnectorClient(connector_base_url)
    local = ZoteroLocalCollection(api_base_url)
    resolver = MetadataResolver()
    collection_key = local.resolve_collection_key(collection_name)
    selected = connector.get_selected_collection()
    target_id = None
    if selected.get("name") == collection_name and selected.get("id"):
        target_id = f"C{selected['id']}" if isinstance(selected.get("id"), int) else selected.get("id")
    else:
        for target in selected.get("targets", []):
            if target.get("name") == collection_name:
                target_id = target.get("id")
                break
    if not target_id:
        raise RuntimeError(f"Could not resolve connector target for collection '{collection_name}'")

    collection_items = local.list_collection_items(collection_key)
    for parsed in references:
        resolution, failure_reason = resolver.resolve(parsed)
        parsed.failure_reason = failure_reason
        if resolution:
            parsed.resolution_source = resolution.source
            parsed.resolution_score = resolution.score
            parsed.used_fallback = False
            parsed.confidence = resolution.source
        else:
            parsed.resolution_source = None
            parsed.resolution_score = None
            parsed.used_fallback = True
            parsed.confidence = "fallback-parsed"

        candidate_item = build_item_payload(parsed, resolution)
        match = find_best_existing_item(candidate_item, collection_items)
        if match and should_reuse_existing_item(match, candidate_item, resolution):
            parsed.item_key = match["key"]
            parsed.created_new_item = False
            continue

        connector.save_item(candidate_item, target=target_id)
        for _ in range(20):
            time.sleep(0.4)
            collection_items = local.list_collection_items(collection_key)
            match = None
            import_marker = candidate_item["_zwb_import_id"]
            for item in collection_items:
                extra = item["data"].get("extra", "")
                if import_marker in extra:
                    match = item
                    break
            if not match:
                match = find_best_existing_item(candidate_item, collection_items)
            if match:
                parsed.item_key = match["key"]
                parsed.created_new_item = True
                break
        if not parsed.item_key:
            raise RuntimeError(f"Imported item not found in collection after save: {parsed.title}")
    return references


def build_citation_code(local: ZoteroLocalCollection, item_keys: list[str]) -> tuple[str, str]:
    citation_items = []
    previews = []
    for item_key in item_keys:
        item_data = local.fetch_csljson(item_key)
        citation_items.append(
            {
                "uris": [local.build_item_uri(item_key)],
                "itemData": item_data,
            }
        )
        previews.append(item_data)
    payload = {
        "citationID": random_id(),
        "properties": {"noteIndex": 0},
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return json.dumps(payload, ensure_ascii=False), ""


def copy_run_properties(source_run) -> etree._Element | None:
    rpr = source_run.find(f"{{{W_NS}}}rPr")
    return etree.fromstring(etree.tostring(rpr)) if rpr is not None else None


def make_run(text: str | None = None, rpr: etree._Element | None = None, *, fld_char: str | None = None, instr: bool = False) -> etree._Element:
    run = etree.Element(f"{{{W_NS}}}r")
    if rpr is not None:
        run.append(etree.fromstring(etree.tostring(rpr)))
    if fld_char:
        etree.SubElement(run, f"{{{W_NS}}}fldChar", attrib={f"{{{W_NS}}}fldCharType": fld_char})
    elif instr:
        node = etree.SubElement(run, f"{{{W_NS}}}instrText")
        node.set(XML_SPACE, "preserve")
        node.text = text or ""
    else:
        node = etree.SubElement(run, f"{{{W_NS}}}t")
        if text and (text[:1].isspace() or text[-1:].isspace()):
            node.set(XML_SPACE, "preserve")
        node.text = text or ""
    return run


def strip_existing_zotero_citation_fields(root: etree._Element) -> None:
    paragraphs = root.xpath(".//w:p", namespaces={"w": W_NS})
    for para in paragraphs:
        runs = para.xpath("./w:r", namespaces={"w": W_NS})
        i = 0
        while i < len(runs):
            run = runs[i]
            fld = run.xpath("./w:fldChar/@w:fldCharType", namespaces={"w": W_NS})
            if fld != ["begin"]:
                i += 1
                continue

            j = i + 1
            instr_runs = []
            result_runs = []
            state = "instr"
            while j < len(runs):
                current = runs[j]
                current_fld = current.xpath("./w:fldChar/@w:fldCharType", namespaces={"w": W_NS})
                if current_fld == ["separate"]:
                    state = "result"
                elif current_fld == ["end"]:
                    break
                else:
                    if state == "instr":
                        instr_runs.append(current)
                    else:
                        result_runs.append(current)
                j += 1
            if j >= len(runs):
                i += 1
                continue

            instr_text = "".join("".join(r.xpath(".//w:instrText/text()", namespaces={"w": W_NS})) for r in instr_runs)
            if "ADDIN ZOTERO_ITEM CSL_CITATION" not in instr_text:
                i = j + 1
                continue

            display_text = "".join("".join(r.xpath(".//w:t/text()", namespaces={"w": W_NS})) for r in result_runs)
            display_rpr = None
            for result_run in result_runs:
                display_rpr = copy_run_properties(result_run)
                if display_rpr is not None:
                    break
            if display_rpr is None:
                display_rpr = copy_run_properties(run)
            replacement = make_run(display_text, rpr=display_rpr)
            start_index = para.index(runs[i])
            for old in runs[i : j + 1]:
                if old.getparent() is para:
                    para.remove(old)
            para.insert(start_index, replacement)
            runs = para.xpath("./w:r", namespaces={"w": W_NS})
            i = start_index + 1


def replace_document_citations(
    input_docx: str | Path,
    output_docx: str | Path,
    references: list[ParsedReference],
    *,
    style_id: str | None = None,
    locale: str | None = None,
) -> None:
    reference_map = {ref.number: ref for ref in references}
    local = ZoteroLocalCollection()
    injector = WordZoteroInjector(None)
    resolved_style_id, resolved_locale = resolve_style_preferences(input_docx, style_id=style_id, locale=locale)

    with ZipFile(input_docx, "r") as source:
        files = {name: source.read(name) for name in source.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    strip_existing_zotero_citation_fields(root)
    paragraphs = root.xpath(".//w:body/w:p", namespaces={"w": W_NS})
    stop_at = None
    for para in paragraphs:
        para_text = "".join(para.xpath(".//w:t/text()", namespaces={"w": W_NS}))
        if normalize_heading(para_text) == "参考文献":
            stop_at = para
            break

    for para in paragraphs:
        if stop_at is not None and para is stop_at:
            break
        runs = para.xpath("./w:r", namespaces={"w": W_NS})
        for run in list(runs):
            text = "".join(run.xpath(".//w:t/text()", namespaces={"w": W_NS}))
            if not text or not CITATION_TOKEN_RE.fullmatch(text):
                continue
            numbers = expand_citation_numbers(text)
            item_keys = []
            for number in numbers:
                parsed = reference_map[number]
                if not parsed.item_key:
                    raise RuntimeError(f"Reference [{number}] has no Zotero item key")
                item_keys.append(parsed.item_key)
            code_payload, _ = build_citation_code(local, item_keys)
            field_code = f" ADDIN ZOTERO_ITEM CSL_CITATION {code_payload} "
            rpr = copy_run_properties(run)
            parent = run.getparent()
            index = parent.index(run)
            new_runs = [
                make_run(rpr=rpr, fld_char="begin"),
                make_run(field_code, rpr=rpr, instr=True),
                make_run(rpr=rpr, fld_char="separate"),
                make_run(text, rpr=rpr),
                make_run(rpr=rpr, fld_char="end"),
            ]
            parent.remove(run)
            for offset, new_run in enumerate(new_runs):
                parent.insert(index + offset, new_run)

    files["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    document_data = injector._build_document_data(style_id=resolved_style_id, locale=resolved_locale, note_type=0)
    files = injector._upsert_custom_properties(files, document_data)

    output_path = Path(output_docx)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as target:
        for name, data in files.items():
            target.writestr(name, data)


def write_report(references: list[ParsedReference], path: str | Path) -> None:
    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(
        json.dumps([asdict(ref) for ref in references], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def write_failure_refs(references: list[ParsedReference], path: str | Path) -> None:
    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    lines = []
    for ref in references:
        if not ref.used_fallback:
            continue
        lines.extend(
            [
                f"[{ref.number}]",
                ref.raw,
                f"failure_reason: {ref.failure_reason or 'no high-confidence metadata source'}",
                f"item_key: {ref.item_key or 'N/A'}",
                "",
            ]
        )
    output.write_text("\n".join(lines), encoding="utf-8")
